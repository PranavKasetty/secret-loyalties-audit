"""Render REPORT.md (and SUPPLEMENT.md) into the organisers' .docx template.

The template carries fonts, margins, heading styles and a theme we do not want
to reproduce by hand. This keeps all of that and replaces only the body: it
opens the template, strips the placeholder paragraphs, and re-emits the
markdown using the template's own named styles. Nothing about the look is
invented here — every paragraph is assigned a style that already exists in the
document.

    python scripts/md_to_docx.py
    python scripts/md_to_docx.py --supplement       # append SUPPLEMENT.md too

Heading map, matched to how the template itself is laid out (numbered sections
are Heading 2 there, not Heading 1):

    # H1   -> Title            ### H3  -> Heading 3
    ## H2  -> Heading 2        #### H4 -> Heading 4

The first `###` line, when it directly follows the title, is treated as the
subtitle rather than a heading, which is how the report is written.

Inline `**bold**`, `*italic*` and `` `code` `` become real runs. Markdown tables
become real Word tables. Images are embedded at a width that fits the text
column. Anything it cannot map is emitted as plain body text rather than
silently dropped.
"""
import argparse
import os
import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from common import ROOT

TEMPLATE = os.path.join(ROOT, "Copy of Secret Loytalties Hackathon submission template.docx")
REPORT = os.path.join(ROOT, "REPORT.md")
SUPPLEMENT = os.path.join(ROOT, "SUPPLEMENT.md")
OUT = os.path.join(ROOT, "submission.docx")

HEADING = {1: "Title", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
CODE_GREY = RGBColor(0x33, 0x33, 0x33)

# The template hard-sets this on every run rather than relying on a style, so
# we must too. Read from the template at runtime rather than hardcoded, so a
# revised template does not silently leave the output in the old typeface.
BODY_FONT = "Old Standard TT"

# Run-level sizes the template applies to its own headings. The style
# definitions disagree with these (Heading 2's style says 16pt while every
# heading run in the document says 14pt); the runs are what render.
HEADING_PT = {2: 14.0, 3: 13.0, 4: 12.0}

# Fills the template's author grid. One entry per author, newline-separated
# name/affiliation, matching the placeholder shape it replaces.
AUTHORS = ["Pranav Kasetty", "Independent researcher"]

# Sits under the author grid, where the template has no slot of its own.
BYLINE_EXTRA = ["Team: Positive Control", "Track 2 — Detection & Auditing"]


def detect_body_font(doc):
    """Most common explicit run font in the template body — the real default."""
    tally = {}
    for par in doc.paragraphs:
        for run in par.runs:
            if run.font.name:
                tally[run.font.name] = tally.get(run.font.name, 0) + 1
    return max(tally, key=tally.get) if tally else BODY_FONT


def clear_body(doc, keep_title_table=True):
    """Remove the placeholder blocks, keeping section properties and the title
    table.

    sectPr holds page size, margins and headers; deleting it throws away the
    page setup, which is half of what the template is for.

    The template's title and abstract live inside its FIRST TABLE, not in body
    paragraphs — that table carries the borders, shading and spacing that make
    the first page look like the template. Deleting it and emitting a bare
    Title-styled paragraph produces a document with the right styles and the
    wrong appearance, which is exactly the mismatch this preserves against.
    """
    body = doc.element.body
    title_tbl = doc.tables[0]._tbl if (keep_title_table and doc.tables) else None
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        if title_tbl is not None and child is title_tbl:
            continue
        body.remove(child)
    return doc.tables[0] if title_tbl is not None else None


def fill_title_table(tbl, title, abstract_paras):
    """Put the title and abstract into the template's own title table.

    Each cell in that table is [empty spacer, the real styled paragraph, empty
    spacer]. The middle one carries the style and the CENTER alignment; the
    spacers are 'normal' with no alignment and exist for vertical rhythm.
    Reusing paragraphs[0] therefore inherits the spacer's formatting and drops
    both the style and the centring — which is what made the first page look
    unlike the template even though the styles matched.

    So: find the paragraph that actually has text, clone its style and
    alignment onto everything written into that cell, and leave the spacers
    alone.
    """
    def set_cell(cell, blocks):
        model = next((p for p in cell.paragraphs if p.text.strip()),
                     cell.paragraphs[0])
        style, align = model.style, model.alignment

        # Compare the underlying XML, not the wrappers: python-docx builds a
        # fresh Paragraph object on every `cell.paragraphs` access, so `is not`
        # against a previously-fetched wrapper is true even for the same
        # element — which silently deleted the very paragraph we meant to keep.
        for par in cell.paragraphs:
            if par._element is not model._element and par.text.strip():
                par._element.getparent().remove(par._element)
        for r in list(model.runs):
            r._element.getparent().remove(r._element)

        add_runs(model, blocks[0])
        model.style, model.alignment = style, align

        anchor = model
        for extra in blocks[1:]:
            par = cell.add_paragraph(style=style)
            par.alignment = align
            add_runs(par, extra)
            anchor._element.addnext(par._element)
            anchor = par

    def nested_tables(cell):
        """Nested tables inside a cell, de-duplicated by XML identity.

        python-docx reports a horizontally-merged cell once per grid column, so
        the same nested table comes back three times for this template's author
        row. Filling it three times writes the author into three phantom
        copies — which is what put six 'Author name N / Affiliation'
        placeholders on the first page.
        """
        seen, out = set(), []
        for t in cell.tables:
            if id(t._tbl) not in seen:
                seen.add(id(t._tbl))
                out.append(t)
        return out

    def write_cell(cell, lines):
        """Replace a cell's entire text content with `lines`."""
        model = next((p for p in cell.paragraphs if p.text.strip()),
                     cell.paragraphs[0])
        style, align = model.style, model.alignment
        for par in cell.paragraphs:
            if par._element is not model._element and par.text.strip():
                par._element.getparent().remove(par._element)
        for r in list(model.runs):
            r._element.getparent().remove(r._element)
        if lines:
            add_runs(model, lines[0])
        model.style, model.alignment = style, align
        anchor = model
        for extra in lines[1:]:
            par = cell.add_paragraph(style=style)
            par.alignment = align
            add_runs(par, extra)
            anchor._element.addnext(par._element)
            anchor = par

    set_cell(tbl.rows[0].cells[0], [title])

    row1 = tbl.rows[1].cells[0]
    inner = nested_tables(row1)

    # The template nests an author grid (2 rows x 3 columns of "Author name N /
    # Affiliation") and, separately, a 1x1 box holding the abstract's guidance
    # text. Both are tables, so a paragraph-only pass leaves them fully intact.
    def find_grid(tables, depth=0):
        """The author grid sits inside the outer 1x3 wrapper, so a search over
        a cell's DIRECT children never reaches it. Recurse."""
        for t in tables:
            if len(t.rows) > 1:
                return t
            for r in t.rows:
                for c in r.cells:
                    found = find_grid(nested_tables(c), depth + 1)
                    if found is not None:
                        return found
        return None

    author_grid = find_grid(inner)
    abstract_box = next((t for t in inner if len(t.rows) == 1
                         and len(t.columns) == 1), None)

    if author_grid is not None:
        seen, cells = set(), []
        for r in author_grid.rows:
            for c in r.cells:
                if id(c._tc) not in seen:
                    seen.add(id(c._tc))
                    cells.append(c)
        write_cell(cells[0], AUTHORS)
        for spare in cells[1:]:
            write_cell(spare, [""])      # blank the five unused author slots

    if abstract_box is not None:
        write_cell(abstract_box.rows[0].cells[0], abstract_paras)
        # Team and track have no template slot; place them just above the
        # "Abstract" label so they read as part of the byline block.
        label = next((p for p in row1.paragraphs
                      if p.text.strip().lower() == "abstract"), None)
        if label is not None:
            anchor = label._element
            for line in BYLINE_EXTRA:
                par = row1.add_paragraph(style=label.style)
                par.alignment = label.alignment
                add_runs(par, line)
                anchor.addprevious(par._element)
    else:
        set_cell(row1, abstract_paras or ["Abstract"])


LINK_BLUE = RGBColor(0x0B, 0x4F, 0x9E)
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def add_hyperlink(par, text, url=None, anchor=None):
    """A real Word hyperlink — clickable in Word and in the exported PDF.

    python-docx has no API for this, so the w:hyperlink element is built by
    hand. `url` makes an external link (relationship + TargetMode="External");
    `anchor` makes an internal jump to a bookmark, which is what turns a [1]
    citation into a link to its reference entry.
    """
    link = OxmlElement("w:hyperlink")
    if url is not None:
        rid = par.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        link.set(f"{{{R_NS}}}id", rid)
    if anchor is not None:
        link.set(f"{W}anchor", anchor)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    for tag, attr, val in (("w:rFonts", "w:ascii", BODY_FONT),
                           ("w:color", "w:val", "0B4F9E")):
        el = OxmlElement(tag)
        el.set(f"{W}{attr.split(':')[1]}", val)
        props.append(el)
    if tag:  # underline, as a link is expected to look
        u = OxmlElement("w:u")
        u.set(f"{W}val", "single")
        props.append(u)
    run.append(props)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    link.append(run)
    par._p.append(link)


def add_bookmark(par, name):
    """Mark a paragraph as a jump target for internal citation links."""
    bid = str(abs(hash(name)) % 100000)
    start = OxmlElement("w:bookmarkStart")
    start.set(f"{W}id", bid)
    start.set(f"{W}name", name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(f"{W}id", bid)
    par._p.insert(0, start)
    par._p.append(end)


def add_runs(par, text):
    """Emit inline **bold**, *italic* and `code` as real runs.

    Every run gets BODY_FONT explicitly. This is not belt-and-braces: the
    template's `normal` style carries no font of its own and its docDefaults
    specify none either, so the typeface is set on each individual run in the
    template's own XML. A run created without one inherits Word's application
    default instead — which is why output with byte-identical styles still came
    out in the wrong typeface.
    """
    # One pass, alternating between literal text and the marked-up spans.
    for piece in re.split(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`)", text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2])
            r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = CODE_GREY
            continue
        elif piece.startswith("*") and piece.endswith("*"):
            r = par.add_run(piece[1:-1])
            r.italic = True
        else:
            emit_linked_text(par, piece)
            continue
        r.font.name = BODY_FONT


# [label](url), <url>, bare http(s) URLs, and [1]-style citations.
_LINKY = re.compile(
    r"\[([^\]]+)\]\((https?://[^)]+)\)"       # markdown link
    r"|<(https?://[^>]+)>"                    # autolink
    r"|(?<![\w/])(https?://[^\s<>)\]]+)"      # bare url
    r"|\[(\d{1,2})\]")                        # citation


def emit_linked_text(par, text):
    """Split plain text into ordinary runs and real, clickable hyperlinks."""
    pos = 0
    for m in _LINKY.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()])
            r.font.name = BODY_FONT
        label, url, auto, bare, cite = m.groups()
        if cite:
            # Internal jump to the matching entry in the reference list.
            add_hyperlink(par, f"[{cite}]", anchor=f"ref{cite}")
        else:
            target = url or auto or bare
            add_hyperlink(par, label or target, url=target)
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:])
        r.font.name = BODY_FONT


def table_style(doc):
    for name in ("Table Grid", "Light Grid", "Normal Table"):
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def add_table(doc, rows, style):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(re.fullmatch(r":?-{2,}:?", x or "-")
                                         for x in c)]
    if not cells:
        return
    width = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=width)
    if style:
        t.style = style
    for i, row in enumerate(cells):
        for j in range(width):
            cell = t.cell(i, j)
            cell.text = ""
            par = cell.paragraphs[0]
            add_runs(par, row[j] if j < len(row) else "")
            for run in par.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True


def convert(doc, md, seen_title):
    lines = md.split("\n")
    style = table_style(doc)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Fenced code -> monospace block, one paragraph per line.
        if line.lstrip().startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            for b in block:
                p = doc.add_paragraph(style="normal")
                r = p.add_run(b)
                r.font.name = "Consolas"
                r.font.size = Pt(8.5)
                p.paragraph_format.space_after = Pt(0)
            continue

        # Image. The alt text wraps across lines in the source, so join the
        # run of lines up to the closing paren before matching.
        if line.lstrip().startswith("!["):
            joined, j = line.strip(), i
            while ")" not in joined and j + 1 < len(lines):
                j += 1
                joined += " " + lines[j].strip()
            i = j
            line = joined
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            path = os.path.join(ROOT, m.group(2))
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                print(f"   ! image not found, skipped: {m.group(2)}")
            i += 1
            continue

        # Table: a run of consecutive pipe rows.
        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows, style)
            doc.add_paragraph("", style="normal")
            continue

        # Headings.
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 3 and not seen_title["sub"] and seen_title["title"]:
                p = doc.add_paragraph(style="Subtitle")
                add_runs(p, text)
                seen_title["sub"] = True
            else:
                p = doc.add_paragraph(style=HEADING.get(level, "Heading 4"))
                add_runs(p, text)
                # The template sizes its heading RUNS, overriding the style.
                if level in HEADING_PT:
                    for r in p.runs:
                        r.font.size = Pt(HEADING_PT[level])
                if level == 1:
                    seen_title["title"] = True
            i += 1
            continue

        # Horizontal rule.
        if re.fullmatch(r"[-*_]{3,}", line.strip()):
            i += 1
            continue

        # Blockquote.
        if line.lstrip().startswith(">"):
            block = []
            while i < len(lines) and (lines[i].lstrip().startswith(">")
                                      or not lines[i].strip()):
                if not lines[i].strip():
                    if block:
                        break
                    i += 1
                    continue
                block.append(lines[i].lstrip()[1:].strip())
                i += 1
            p = doc.add_paragraph(style="normal")
            p.paragraph_format.left_indent = Inches(0.4)
            add_runs(p, " ".join(block))
            for r in p.runs:
                r.italic = True
            continue

        # List item: gather its wrapped continuation lines.
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(3)
            i += 1
            while (i < len(lines) and lines[i].strip()
                   and not re.match(r"^(\s*)([-*+]|\d+\.)\s+", lines[i])
                   and not lines[i].startswith("#")
                   and not lines[i].strip().startswith("|")
                   and lines[i].startswith(" ")):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="normal")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
            bullet = m.group(2) if m.group(2)[0].isdigit() else "•"
            add_runs(p, f"{bullet}  {text}")
            continue

        # Body paragraph: join wrapped lines until a blank or a block starter.
        block = [line]
        i += 1
        while (i < len(lines) and lines[i].strip()
               and not lines[i].startswith(("#", "|", ">", "```", "!["))
               and not re.match(r"^(\s*)([-*+]|\d+\.)\s+", lines[i])):
            block.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph(style="normal")
        body_text = " ".join(x.strip() for x in block)
        m_ref = re.match(r"^\[(\d{1,2})\]\s", body_text)
        if m_ref:
            # A reference-list entry: bookmark it as a citation target, and do
            # not turn its own leading [n] into a link pointing at itself.
            add_bookmark(p, "ref" + m_ref.group(1))
            r = p.add_run("[" + m_ref.group(1) + "] ")
            r.font.name = BODY_FONT
            r.bold = True
            add_runs(p, body_text[m_ref.end():])
        else:
            add_runs(p, body_text)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--supplement", action="store_true",
                    help="append SUPPLEMENT.md after the report, on a new page")
    ap.add_argument("--out", default=OUT)
    args = ap.parse_args()

    if not os.path.exists(TEMPLATE):
        raise SystemExit(f"template not found: {TEMPLATE}")
    if not os.path.exists(REPORT):
        raise SystemExit("REPORT.md not found — run scripts/make_report.py first")

    doc = docx.Document(TEMPLATE)
    global BODY_FONT
    BODY_FONT = detect_body_font(doc)
    print(f"  body font from template: {BODY_FONT}")
    tbl = clear_body(doc)

    md = open(REPORT, encoding="utf-8").read()

    # Split off the title and abstract; they belong in the template's table.
    NL = chr(10)
    title = md.split(NL, 1)[0].lstrip("# ").strip()
    abs_start = md.index("## Abstract") + len("## Abstract")
    abs_end = md.index("## 1. Introduction")
    abstract = [b.replace(NL, " ").strip()
                for b in md[abs_start:abs_end].strip().split(NL + NL) if b.strip()]
    byline = [l.strip() for l in md.split("## Abstract")[0].split(NL)[1:]
              if l.strip() and not l.startswith("#")]

    if tbl is not None:
        fill_title_table(tbl, title, abstract)
        body_md = md[abs_end:]
    else:
        body_md = md

    seen = {"title": True, "sub": True}
    convert(doc, body_md, seen)

    if args.supplement and os.path.exists(SUPPLEMENT):
        doc.add_page_break()
        convert(doc, open(SUPPLEMENT, encoding="utf-8").read(), seen)

    doc.save(args.out)
    n_par = len(doc.paragraphs)
    print(f"wrote {args.out}")
    print(f"  {n_par} paragraphs, {len(doc.tables)} tables, "
          f"{len(doc.inline_shapes)} images")
    print("  Open it, check the first page, then File > Save as PDF.")


if __name__ == "__main__":
    main()
